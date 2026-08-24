import os
import pandas as pd
from pathlib import Path
from pdf2docx import Converter
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
from docx import Document
import warnings

# Suppress ebooklib warnings about missing nav elements
warnings.filterwarnings('ignore', category=UserWarning, module='ebooklib')

def epub_to_docx(epub_path, docx_path):
    """Converts an EPUB file to a DOCX file by extracting text and basic headings."""
    book = epub.read_epub(str(epub_path))
    doc = Document()
    
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            soup = BeautifulSoup(item.get_body_content(), 'html.parser')
            for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                text = element.get_text(strip=True)
                if text:
                    if element.name.startswith('h'):
                        try:
                            level = int(element.name[1])
                            doc.add_heading(text, level=min(level, 9)) # docx only supports up to level 9
                        except Exception:
                            doc.add_paragraph(text)
                    else:
                        doc.add_paragraph(text)
                        
    doc.save(str(docx_path))

def process_books(excel_path: str, input_dir: str):
    """
    Reads an Excel file containing book details and converts corresponding PDFs or EPUBs to Word.
    Organizes the output in the user's Downloads folder by Series Name.
    
    Expected Excel Columns:
    - 'Book Title': The name of the book
    - 'Series Name': The series it belongs to (used for folder creation)
    - 'Filename': The exact filename of the downloaded file (e.g., 'book1.pdf' or 'book2.epub')
    """
    downloads_folder = Path.home() / "Downloads"
    
    print(f"Reading Excel file from: {excel_path}")
    try:
        df = pd.read_excel(excel_path)
    except FileNotFoundError:
        print(f"Error: Could not find the Excel file at {excel_path}")
        return
    except Exception as e:
        print(f"Error reading Excel file: {e}")
        return

    # Ensure required columns exist
    required_columns = ['Book Title', 'Series Name', 'Filename']
    if not all(col in df.columns for col in required_columns):
        print(f"Error: Excel file must contain the following columns: {required_columns}")
        # Backwards compatibility check: If they still have 'PDF Filename', rename it.
        if 'PDF Filename' in df.columns:
             df.rename(columns={'PDF Filename': 'Filename'}, inplace=True)
             print("Found 'PDF Filename' column, interpreting it as 'Filename'.")
        else:
             return

    for index, row in df.iterrows():
        book_title = row['Book Title']
        series_name = str(row['Series Name']) if pd.notna(row['Series Name']) else "Standalone Books"
        filename = str(row['Filename'])

        file_path = Path(input_dir) / filename

        if not file_path.exists():
            print(f"[{book_title}] File not found at {file_path}. Skipping.")
            continue

        # Create Series folder in Downloads
        series_folder = downloads_folder / series_name
        series_folder.mkdir(parents=True, exist_ok=True)

        # Output Word file path (.docx)
        docx_filename = file_path.stem + ".docx"
        docx_path = series_folder / docx_filename
        
        if docx_path.exists():
            print(f"[{book_title}] Word file already exists at {docx_path}. Skipping.")
            continue

        file_ext = file_path.suffix.lower()

        try:
            if file_ext == '.pdf':
                print(f"[{book_title}] Converting PDF '{filename}' to Word...")
                cv = Converter(str(file_path))
                cv.convert(str(docx_path))
                cv.close()
                print(f"[{book_title}] Successfully saved to: {docx_path}\n")
                
            elif file_ext == '.epub':
                print(f"[{book_title}] Converting EPUB '{filename}' to Word...")
                epub_to_docx(file_path, docx_path)
                print(f"[{book_title}] Successfully saved to: {docx_path}\n")
                
            else:
                print(f"[{book_title}] Unsupported file format '{file_ext}'. Only .pdf and .epub are supported.\n")
                
        except Exception as e:
            print(f"[{book_title}] Error during conversion: {e}\n")

if __name__ == "__main__":
    # --- CONFIGURATION ---
    EXCEL_FILE_PATH = r"path\to\your\excel\sheet.xlsx" 
    DOWNLOAD_DIRECTORY = r"path\to\where\you\saved\files"
    
    print("Script started. Please ensure you have updated the paths in the script before running.")
    # Uncomment the line below to actually run the script once paths are set
    # process_books(EXCEL_FILE_PATH, DOWNLOAD_DIRECTORY)
